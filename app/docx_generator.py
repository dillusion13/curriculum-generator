"""
DOCX Generator - Create editable Word documents from curriculum JSON.
Produces a single combined document with Teacher Guide and all Student Materials.
"""
from pathlib import Path
from typing import Any, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

from .docx_styles import (
    COLORS, FONTS, FONT_SIZES, SPACING,
    get_color, hex_to_rgb, get_level_accent, get_level_light, get_level_name,
    set_cell_shading, set_cell_border, set_cell_borders, set_table_borders,
    remove_cell_borders,
)


# Level display names
LEVEL_NAMES = {
    "below_level": "Below Level",
    "approaching_level": "Approaching Level",
    "at_level": "At Level",
    "above_level": "Above Level",
}


def _add_page_numbers(doc: Document) -> None:
    """Add page numbers to the document footer.

    Creates a centered footer with "Page X of Y" format.
    """
    section = doc.sections[0]

    # Create footer
    footer = section.footer
    footer.is_linked_to_previous = False

    # Get or create paragraph in footer
    if footer.paragraphs:
        para = footer.paragraphs[0]
    else:
        para = footer.add_paragraph()

    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add "Page " text
    run = para.add_run("Page ")
    run.font.name = FONTS["body"]
    run.font.size = Pt(9)
    run.font.color.rgb = get_color("ink_500")

    # Add PAGE field (current page number) - must be wrapped in runs
    run_begin = para.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run_begin._r.append(fldChar1)

    run_instr = para.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run_instr._r.append(instrText)

    run_sep = para.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run_sep._r.append(fldChar2)

    run_end = para.add_run()
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run_end._r.append(fldChar3)

    # Add " of " text
    run2 = para.add_run(" of ")
    run2.font.name = FONTS["body"]
    run2.font.size = Pt(9)
    run2.font.color.rgb = get_color("ink_500")

    # Add NUMPAGES field (total pages) - must be wrapped in runs
    run_begin2 = para.add_run()
    fldChar4 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run_begin2._r.append(fldChar4)

    run_instr2 = para.add_run()
    instrText2 = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> NUMPAGES </w:instrText>')
    run_instr2._r.append(instrText2)

    run_sep2 = para.add_run()
    fldChar5 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run_sep2._r.append(fldChar5)

    run_end2 = para.add_run()
    fldChar6 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run_end2._r.append(fldChar6)


def _add_styled_heading(doc: Document, text: str, level: int = 1, accent_color: RGBColor = None) -> None:
    """Add a styled heading to the document.

    Args:
        doc: The document to add to
        text: Heading text
        level: Heading level (1=title, 2=section)
        accent_color: Optional accent color for level 2 headings
    """
    heading = doc.add_heading(text, level=level)
    run = heading.runs[0]

    if level == 1:
        run.font.name = FONTS["heading"]
        run.font.size = FONT_SIZES["title"]
        run.font.color.rgb = get_color("ink_900")
        run.font.bold = True
    elif level == 2:
        run.font.name = FONTS["heading"]
        run.font.size = FONT_SIZES["heading1"]
        run.font.color.rgb = accent_color if accent_color else get_color("ink_800")
        run.font.bold = True


def _add_section_header(doc: Document, title: str, accent_color: RGBColor = None) -> None:
    """Add a section header with colored left border.

    Creates a table-based header with background and accent border.
    """
    if accent_color is None:
        accent_color = get_color("navy_700")

    # Create a single-cell table for the header
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.allow_autofit = False

    # Set table width to full page
    table.columns[0].width = Inches(7.0)

    cell = table.rows[0].cells[0]
    cell.width = Inches(7.0)

    # Add the header text
    para = cell.paragraphs[0]
    run = para.add_run(title.upper())
    run.bold = True
    run.font.name = FONTS["body"]
    run.font.size = FONT_SIZES["heading2"]
    run.font.color.rgb = get_color("ink_800")
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)

    # Style the cell: gray background + accent left border
    set_cell_shading(cell, COLORS["ink_100"])

    # Get hex color for border
    accent_hex = f"{accent_color.red:02x}{accent_color.green:02x}{accent_color.blue:02x}" if hasattr(accent_color, 'red') else COLORS["navy_700"]
    set_cell_border(cell, "left", accent_hex, width=24, style="single")  # 3pt left border
    set_cell_border(cell, "top", COLORS["ink_200"], width=0, style="nil")
    set_cell_border(cell, "right", COLORS["ink_200"], width=0, style="nil")
    set_cell_border(cell, "bottom", COLORS["ink_200"], width=8, style="single")

    # Add spacing after the header
    doc.add_paragraph()


def _add_key_value(doc: Document, key: str, value: str) -> None:
    """Add a key-value pair with consistent typography."""
    para = doc.add_paragraph()
    key_run = para.add_run(f"{key}: ")
    key_run.bold = True
    key_run.font.name = FONTS["body"]
    key_run.font.size = FONT_SIZES["body"]
    key_run.font.color.rgb = get_color("ink_800")

    value_run = para.add_run(str(value))
    value_run.font.name = FONTS["body"]
    value_run.font.size = FONT_SIZES["body"]
    value_run.font.color.rgb = get_color("ink_700")


def _add_bullet_list(doc: Document, items: list, accent_color: RGBColor = None) -> None:
    """Add a bullet list with consistent typography."""
    for item in items:
        para = doc.add_paragraph(style='List Bullet')
        run = para.add_run(str(item))
        run.font.name = FONTS["body"]
        run.font.size = FONT_SIZES["body"]
        run.font.color.rgb = get_color("ink_700")
        para.paragraph_format.space_after = SPACING["list_item_after"]


def _add_numbered_list(doc: Document, items: list) -> None:
    """Add a numbered list with consistent typography."""
    for item in items:
        para = doc.add_paragraph(style='List Number')
        run = para.add_run(str(item))
        run.font.name = FONTS["body"]
        run.font.size = FONT_SIZES["body"]
        run.font.color.rgb = get_color("ink_700")
        para.paragraph_format.space_after = SPACING["list_item_after"]


def _add_table(doc: Document, headers: list, rows: list, accent_color: RGBColor = None) -> None:
    """Add a styled table with headers and rows.

    Args:
        doc: Document to add table to
        headers: List of header strings
        rows: List of row data (each row is a list of cell values)
        accent_color: Optional accent color for header row
    """
    if not rows:
        return

    if accent_color is None:
        accent_color = get_color("navy_700")

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    # Get hex color for header background
    accent_hex = f"{accent_color.red:02x}{accent_color.green:02x}{accent_color.blue:02x}" if hasattr(accent_color, 'red') else COLORS["navy_700"]

    # Style header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        cell = header_cells[i]
        cell.text = header
        para = cell.paragraphs[0]
        for run in para.runs:
            run.bold = True
            run.font.name = FONTS["body"]
            run.font.size = FONT_SIZES["body"]
            run.font.color.rgb = get_color("white")
        set_cell_shading(cell, accent_hex)

    # Style data rows
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_data in enumerate(row_data):
            cell = row_cells[col_idx]
            cell.text = str(cell_data) if cell_data else ""
            para = cell.paragraphs[0]
            for run in para.runs:
                run.font.name = FONTS["body"]
                run.font.size = FONT_SIZES["body"]
                run.font.color.rgb = get_color("ink_700")

    doc.add_paragraph()


def _add_info_box(doc: Document, title: str, content: str, accent_color: RGBColor = None) -> None:
    """Add an info box with colored left border."""
    if accent_color is None:
        accent_color = get_color("navy_700")

    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(7.0)

    cell = table.rows[0].cells[0]
    cell.width = Inches(7.0)

    para = cell.paragraphs[0]
    title_run = para.add_run(f"{title}: ")
    title_run.bold = True
    title_run.font.name = FONTS["body"]
    title_run.font.size = FONT_SIZES["body"]
    title_run.font.color.rgb = get_color("ink_800")

    content_run = para.add_run(content)
    content_run.font.name = FONTS["body"]
    content_run.font.size = FONT_SIZES["body"]
    content_run.font.color.rgb = get_color("ink_700")

    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)

    # Style: light background + accent left border
    set_cell_shading(cell, COLORS["ink_50"])
    accent_hex = f"{accent_color.red:02x}{accent_color.green:02x}{accent_color.blue:02x}" if hasattr(accent_color, 'red') else COLORS["navy_700"]
    set_cell_border(cell, "left", accent_hex, width=24, style="single")
    set_cell_border(cell, "top", COLORS["ink_200"], width=8, style="single")
    set_cell_border(cell, "right", COLORS["ink_200"], width=8, style="single")
    set_cell_border(cell, "bottom", COLORS["ink_200"], width=8, style="single")

    doc.add_paragraph()


def _add_quick_reference_box(doc: Document, meta: dict) -> None:
    """Add a compact quick reference box with key lesson info.

    ┌─────────────────────────────────────────────────┐
    │ QUICK REFERENCE                                 │
    │ Duration: 45 min │ Format: Small Group          │
    │ Approach: 5E Lesson │ Standard: 7.NS.A.1        │
    └─────────────────────────────────────────────────┘
    """
    # Gather info
    duration = meta.get("duration_minutes", "")
    approach_data = meta.get("pedagogical_approach", "")
    if isinstance(approach_data, dict):
        approach = approach_data.get("name", approach_data.get("id", ""))
    else:
        approach = str(approach_data) if approach_data else ""
    grouping = meta.get("grouping", meta.get("format", ""))
    standards = meta.get("standards_addressed", [])
    key_standard = standards[0] if standards else ""

    # Build info items
    info_items = []
    if duration:
        info_items.append(("Duration", f"{duration} min"))
    if grouping:
        info_items.append(("Format", grouping))
    if approach:
        info_items.append(("Approach", approach))
    if key_standard:
        info_items.append(("Standard", key_standard))

    if not info_items:
        return

    # Create table
    table = doc.add_table(rows=2, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(7.0)

    # Header row
    header_cell = table.rows[0].cells[0]
    header_para = header_cell.paragraphs[0]
    header_run = header_para.add_run("QUICK REFERENCE")
    header_run.bold = True
    header_run.font.name = FONTS["body"]
    header_run.font.size = FONT_SIZES["small"]
    header_run.font.color.rgb = get_color("navy_700")
    set_cell_shading(header_cell, COLORS["navy_100"])
    set_cell_border(header_cell, "top", COLORS["navy_700"], width=8)
    set_cell_border(header_cell, "left", COLORS["navy_700"], width=8)
    set_cell_border(header_cell, "right", COLORS["navy_700"], width=8)
    set_cell_border(header_cell, "bottom", COLORS["ink_200"], width=4)

    # Content row
    content_cell = table.rows[1].cells[0]
    content_para = content_cell.paragraphs[0]

    for i, (key, value) in enumerate(info_items):
        if i > 0:
            sep_run = content_para.add_run("  │  ")
            sep_run.font.name = FONTS["body"]
            sep_run.font.size = FONT_SIZES["small"]
            sep_run.font.color.rgb = get_color("ink_300")

        key_run = content_para.add_run(f"{key}: ")
        key_run.bold = True
        key_run.font.name = FONTS["body"]
        key_run.font.size = FONT_SIZES["small"]
        key_run.font.color.rgb = get_color("ink_600")

        value_run = content_para.add_run(value)
        value_run.font.name = FONTS["body"]
        value_run.font.size = FONT_SIZES["small"]
        value_run.font.color.rgb = get_color("ink_700")

    set_cell_shading(content_cell, COLORS["navy_100"])
    set_cell_border(content_cell, "top", COLORS["ink_200"], width=0, style="nil")
    set_cell_border(content_cell, "left", COLORS["navy_700"], width=8)
    set_cell_border(content_cell, "right", COLORS["navy_700"], width=8)
    set_cell_border(content_cell, "bottom", COLORS["navy_700"], width=8)

    doc.add_paragraph()


def _add_differentiation_at_a_glance(doc: Document, diff: dict) -> None:
    """Add a 4-column differentiation summary table with level colors."""
    level_info = [
        ("below_level", "Below", COLORS["below"], COLORS["below_light"]),
        ("approaching_level", "Approaching", COLORS["approaching"], COLORS["approaching_light"]),
        ("at_level", "At Level", COLORS["at"], COLORS["at_light"]),
        ("above_level", "Above", COLORS["above"], COLORS["above_light"]),
    ]

    # Create 4-column table (header + content)
    table = doc.add_table(rows=2, cols=4)
    table.autofit = False
    col_width = Inches(1.75)
    for col in table.columns:
        col.width = col_width

    # Header row with level names
    for i, (key, name, accent_hex, light_hex) in enumerate(level_info):
        cell = table.rows[0].cells[i]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = para.add_run(name)
        run.bold = True
        run.font.name = FONTS["body"]
        run.font.size = FONT_SIZES["small"]
        run.font.color.rgb = hex_to_rgb(accent_hex)

        set_cell_shading(cell, light_hex)
        set_cell_border(cell, "top", accent_hex, width=8)
        set_cell_border(cell, "left", COLORS["ink_200"], width=4)
        set_cell_border(cell, "right", COLORS["ink_200"], width=4)
        set_cell_border(cell, "bottom", COLORS["ink_200"], width=4)

    # Content row with focus for each level
    for i, (key, name, accent_hex, light_hex) in enumerate(level_info):
        cell = table.rows[1].cells[i]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        level_data = diff.get(key, {})
        focus = level_data.get("focus", "") if isinstance(level_data, dict) else ""

        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(focus)
        run.font.name = FONTS["body"]
        run.font.size = FONT_SIZES["small"]
        run.font.color.rgb = get_color("ink_700")

        set_cell_shading(cell, light_hex)
        set_cell_border(cell, "top", COLORS["ink_200"], width=0, style="nil")
        set_cell_border(cell, "left", COLORS["ink_200"], width=4)
        set_cell_border(cell, "right", COLORS["ink_200"], width=4)
        set_cell_border(cell, "bottom", accent_hex, width=8)

    doc.add_paragraph()


def _add_styled_paragraph(doc: Document, text: str, bold: bool = False, italic: bool = False, color_key: str = "ink_700") -> None:
    """Add a paragraph with consistent font styling."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = FONTS["body"]
    run.font.size = FONT_SIZES["body"]
    run.font.color.rgb = get_color(color_key)
    run.bold = bold
    run.italic = italic
    return para


def _add_workspace_box(doc: Document, num_lines: int = 4, accent_color: RGBColor = None) -> None:
    """Add a workspace box with dotted writing lines.

    Args:
        doc: Document to add to
        num_lines: Number of writing lines
        accent_color: Color for left border accent
    """
    if accent_color is None:
        accent_color = get_color("ink_400")

    table = doc.add_table(rows=num_lines, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(7.0)

    accent_hex = f"{accent_color.red:02x}{accent_color.green:02x}{accent_color.blue:02x}" if hasattr(accent_color, 'red') else COLORS["ink_400"]

    for i, row in enumerate(table.rows):
        row.height = Twips(400)  # ~0.28 inch per line
        cell = row.cells[0]
        cell.width = Inches(7.0)

        # Add empty paragraph with minimum height
        para = cell.paragraphs[0]
        para.paragraph_format.space_after = Pt(0)

        # Left accent border for first cell only
        if i == 0:
            set_cell_border(cell, "left", accent_hex, width=24, style="single")
            set_cell_border(cell, "top", COLORS["ink_200"], width=8, style="single")
        else:
            set_cell_border(cell, "left", accent_hex, width=24, style="single")
            set_cell_border(cell, "top", COLORS["ink_200"], width=0, style="nil")

        # Dotted bottom border for writing lines
        set_cell_border(cell, "bottom", COLORS["ink_300"], width=4, style="dotted")
        set_cell_border(cell, "right", COLORS["ink_200"], width=8, style="single")

    doc.add_paragraph()


def _add_goal_box(doc: Document, i_can_statement: str, accent_color: RGBColor = None) -> None:
    """Add a colored goal box with I CAN statement.

    Args:
        doc: Document to add to
        i_can_statement: The I CAN statement text
        accent_color: Accent color for text and border
    """
    if accent_color is None:
        accent_color = get_color("navy_700")

    # Get corresponding light background color
    accent_hex = f"{accent_color.red:02x}{accent_color.green:02x}{accent_color.blue:02x}" if hasattr(accent_color, 'red') else COLORS["navy_700"]

    # Determine light background based on accent color (approximate matching)
    light_bg = COLORS["ink_100"]  # default
    if accent_hex.lower() == COLORS["below"].lower():
        light_bg = COLORS["below_light"]
    elif accent_hex.lower() == COLORS["approaching"].lower():
        light_bg = COLORS["approaching_light"]
    elif accent_hex.lower() == COLORS["at"].lower():
        light_bg = COLORS["at_light"]
    elif accent_hex.lower() == COLORS["above"].lower():
        light_bg = COLORS["above_light"]
    elif accent_hex.lower() == COLORS["navy_700"].lower():
        light_bg = COLORS["navy_100"]

    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(7.0)

    cell = table.rows[0].cells[0]
    cell.width = Inches(7.0)

    para = cell.paragraphs[0]
    # Add goal label
    label_run = para.add_run("TODAY'S GOAL: ")
    label_run.bold = True
    label_run.font.name = FONTS["body"]
    label_run.font.size = FONT_SIZES["i_can"]
    label_run.font.color.rgb = accent_color

    # Add the I CAN statement
    statement_run = para.add_run(i_can_statement)
    statement_run.font.name = FONTS["body"]
    statement_run.font.size = FONT_SIZES["i_can"]
    statement_run.font.color.rgb = accent_color
    statement_run.italic = True

    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(8)

    # Style: light background + accent left border
    set_cell_shading(cell, light_bg)
    set_cell_border(cell, "left", accent_hex, width=24, style="single")
    set_cell_border(cell, "top", light_bg, width=0, style="nil")
    set_cell_border(cell, "right", light_bg, width=0, style="nil")
    set_cell_border(cell, "bottom", light_bg, width=0, style="nil")

    doc.add_paragraph()


def _add_student_header(doc: Document, title: str, level_name: str, accent_color: RGBColor = None) -> None:
    """Add student material header with title, name field, and date field.

    Args:
        doc: Document to add to
        title: Lesson title
        level_name: Display name for the readiness level
        accent_color: Accent color for the level
    """
    if accent_color is None:
        accent_color = get_color("ink_700")

    accent_hex = f"{accent_color.red:02x}{accent_color.green:02x}{accent_color.blue:02x}" if hasattr(accent_color, 'red') else COLORS["ink_700"]

    # Create header table: Title | Name/Date fields
    table = doc.add_table(rows=2, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(5.0)
    table.columns[1].width = Inches(2.0)

    # Row 1: Title and Name field
    title_cell = table.rows[0].cells[0]
    title_para = title_cell.paragraphs[0]
    title_run = title_para.add_run(title)
    title_run.font.name = FONTS["heading"]
    title_run.font.size = FONT_SIZES["title"]
    title_run.font.color.rgb = get_color("ink_900")
    title_run.bold = True

    name_cell = table.rows[0].cells[1]
    name_para = name_cell.paragraphs[0]
    name_run = name_para.add_run("Name: ____________________")
    name_run.font.name = FONTS["body"]
    name_run.font.size = FONT_SIZES["body"]
    name_run.font.color.rgb = get_color("ink_600")
    name_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Row 2: Level indicator and Date field
    level_cell = table.rows[1].cells[0]
    level_para = level_cell.paragraphs[0]
    level_run = level_para.add_run(level_name)
    level_run.font.name = FONTS["body"]
    level_run.font.size = FONT_SIZES["small"]
    level_run.font.color.rgb = accent_color
    level_run.bold = True

    date_cell = table.rows[1].cells[1]
    date_para = date_cell.paragraphs[0]
    date_run = date_para.add_run("Date: __________")
    date_run.font.name = FONTS["body"]
    date_run.font.size = FONT_SIZES["body"]
    date_run.font.color.rgb = get_color("ink_600")
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Remove all borders
    for row in table.rows:
        for cell in row.cells:
            remove_cell_borders(cell)

    # Add accent bar below header
    bar_table = doc.add_table(rows=1, cols=1)
    bar_table.autofit = False
    bar_table.columns[0].width = Inches(7.0)
    bar_cell = bar_table.rows[0].cells[0]
    bar_cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    # Make it just an accent line
    set_cell_shading(bar_cell, accent_hex)
    bar_table.rows[0].height = Twips(80)  # Thin accent bar
    remove_cell_borders(bar_cell)

    doc.add_paragraph()


def _render_graphic_organizer(doc: Document, organizer: dict, accent_color: RGBColor = None) -> None:
    """Render a graphic organizer as an actual table.

    Supports: ratio_table, table, t_chart, comparison, vocabulary_four_square, four_square
    """
    if not organizer:
        return

    if accent_color is None:
        accent_color = get_color("navy_700")

    org_type = organizer.get("type", "").lower()
    title = organizer.get("title", "")
    description = organizer.get("description", "")

    accent_hex = f"{accent_color.red:02x}{accent_color.green:02x}{accent_color.blue:02x}" if hasattr(accent_color, 'red') else COLORS["navy_700"]

    # Add title if present
    if title:
        para = doc.add_paragraph()
        run = para.add_run(title)
        run.bold = True
        run.font.name = FONTS["body"]
        run.font.size = FONT_SIZES["body"]
        run.font.color.rgb = get_color("ink_800")

    # Render based on type
    if org_type in ["ratio_table", "table", "data_table"]:
        # Table organizer
        headers = organizer.get("headers", organizer.get("columns", []))
        rows = organizer.get("rows", organizer.get("data", []))

        if headers:
            num_cols = len(headers)
            num_rows = len(rows) if rows else 4  # Default 4 empty rows

            table = doc.add_table(rows=1 + num_rows, cols=num_cols)
            table.style = 'Table Grid'

            # Header row
            for i, header in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = str(header)
                for run in cell.paragraphs[0].runs:
                    run.bold = True
                    run.font.name = FONTS["body"]
                    run.font.size = FONT_SIZES["body"]
                    run.font.color.rgb = get_color("white")
                set_cell_shading(cell, accent_hex)

            # Data rows
            if rows:
                for row_idx, row_data in enumerate(rows):
                    for col_idx, cell_data in enumerate(row_data if isinstance(row_data, list) else [row_data]):
                        if col_idx < num_cols:
                            cell = table.rows[row_idx + 1].cells[col_idx]
                            cell.text = str(cell_data) if cell_data else ""

    elif org_type in ["t_chart", "comparison", "t-chart"]:
        # T-Chart: two columns
        left_label = organizer.get("left_label", organizer.get("side_a", "Side A"))
        right_label = organizer.get("right_label", organizer.get("side_b", "Side B"))
        num_rows = organizer.get("rows", 4)

        table = doc.add_table(rows=1 + num_rows, cols=2)
        table.style = 'Table Grid'

        # Headers
        for i, label in enumerate([left_label, right_label]):
            cell = table.rows[0].cells[i]
            cell.text = str(label)
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.name = FONTS["body"]
                run.font.size = FONT_SIZES["body"]
                run.font.color.rgb = get_color("white")
            set_cell_shading(cell, accent_hex)

        # Empty rows for student work
        for row_idx in range(1, num_rows + 1):
            for cell in table.rows[row_idx].cells:
                cell.paragraphs[0].paragraph_format.space_after = Pt(20)

    elif org_type in ["vocabulary_four_square", "four_square", "4_square"]:
        # 2x2 grid with center term
        term = organizer.get("term", organizer.get("word", ""))
        quadrants = ["Definition", "Example", "Non-Example", "Picture/Drawing"]

        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'

        # Top row quadrants
        for i, label in enumerate(quadrants[:2]):
            cell = table.rows[0].cells[i]
            para = cell.paragraphs[0]
            label_run = para.add_run(f"{label}:\n")
            label_run.bold = True
            label_run.font.name = FONTS["body"]
            label_run.font.size = FONT_SIZES["small"]
            label_run.font.color.rgb = get_color("ink_600")
            # Add space for writing
            para.add_run("\n\n")

        # Middle row: center term
        center_cell = table.rows[1].cells[0]
        table.rows[1].cells[0].merge(table.rows[1].cells[1])
        center_para = center_cell.paragraphs[0]
        center_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        term_run = center_para.add_run(term)
        term_run.bold = True
        term_run.font.name = FONTS["heading"]
        term_run.font.size = FONT_SIZES["heading1"]
        term_run.font.color.rgb = get_color("white")
        set_cell_shading(center_cell, accent_hex)

        # Bottom row quadrants
        for i, label in enumerate(quadrants[2:]):
            cell = table.rows[2].cells[i]
            para = cell.paragraphs[0]
            label_run = para.add_run(f"{label}:\n")
            label_run.bold = True
            label_run.font.name = FONTS["body"]
            label_run.font.size = FONT_SIZES["small"]
            label_run.font.color.rgb = get_color("ink_600")
            para.add_run("\n\n")

    else:
        # Fallback: show description and workspace
        if description:
            para = doc.add_paragraph()
            run = para.add_run(description)
            run.font.name = FONTS["body"]
            run.font.size = FONT_SIZES["body"]
            run.font.color.rgb = get_color("ink_700")

        _add_workspace_box(doc, num_lines=4, accent_color=accent_color)

    doc.add_paragraph()


def generate_teacher_guide_section(doc: Document, teacher_guide: dict, day_num: Optional[int] = None) -> None:
    """Generate teacher guide section in the document."""
    meta = teacher_guide.get("metadata", {})
    navy_accent = get_color("navy_700")

    # Title
    title = meta.get("title", "Lesson Plan")
    if day_num:
        title = f"Day {day_num}: {title}"
    _add_styled_heading(doc, title, level=1)

    # Quick Reference Box (new)
    _add_quick_reference_box(doc, meta)

    # Metadata section
    _add_section_header(doc, "Lesson Overview", navy_accent)
    _add_key_value(doc, "Grade", meta.get("grade", ""))
    _add_key_value(doc, "Subject", meta.get("subject", ""))
    _add_key_value(doc, "Topic", meta.get("topic", ""))
    _add_key_value(doc, "Duration", f"{meta.get('duration_minutes', '')} minutes")

    standards = meta.get("standards_addressed", [])
    if standards:
        _add_key_value(doc, "Standards", ", ".join(standards))

    approach = meta.get("pedagogical_approach", {})
    if isinstance(approach, dict):
        _add_key_value(doc, "Approach", approach.get("name", ""))
        if approach.get("rationale"):
            _add_styled_paragraph(doc, f"Rationale: {approach['rationale']}", italic=True, color_key="ink_600")

    doc.add_paragraph()

    # Learning Objectives
    objectives = teacher_guide.get("learning_objectives", [])
    if objectives:
        _add_section_header(doc, "Learning Objectives", navy_accent)
        for obj in objectives:
            if isinstance(obj, dict):
                para = doc.add_paragraph()
                obj_run = para.add_run(obj.get("objective", ""))
                obj_run.bold = True
                obj_run.font.name = FONTS["body"]
                obj_run.font.size = FONT_SIZES["body"]
                obj_run.font.color.rgb = get_color("ink_800")
                if obj.get("success_criteria"):
                    _add_bullet_list(doc, [f"Success Criteria: {obj['success_criteria']}"])
            else:
                _add_bullet_list(doc, [str(obj)])
        doc.add_paragraph()

    # Differentiation at a Glance (new)
    diff = teacher_guide.get("differentiation_overview", {})
    if diff:
        _add_section_header(doc, "Differentiation at a Glance", navy_accent)
        _add_differentiation_at_a_glance(doc, diff)

    # Session Structure / Phases
    session = teacher_guide.get("session_structure", {})
    phases = session.get("phases", [])
    if phases:
        _add_section_header(doc, "Session Structure", navy_accent)
        for i, phase in enumerate(phases, 1):
            # Phase header with styling
            para = doc.add_paragraph()
            phase_name = phase.get("name", f"Phase {i}")
            duration = phase.get("duration_minutes", "")
            name_run = para.add_run(f"{phase_name}")
            name_run.bold = True
            name_run.font.name = FONTS["body"]
            name_run.font.size = FONT_SIZES["heading2"]
            name_run.font.color.rgb = get_color("navy_700")
            if duration:
                dur_run = para.add_run(f" ({duration} min)")
                dur_run.font.name = FONTS["body"]
                dur_run.font.size = FONT_SIZES["body"]
                dur_run.font.color.rgb = get_color("ink_500")

            # Phase details with styled paragraphs
            if phase.get("description"):
                _add_styled_paragraph(doc, phase["description"])
            if phase.get("teacher_actions"):
                para = doc.add_paragraph()
                label = para.add_run("Teacher Actions: ")
                label.bold = True
                label.font.name = FONTS["body"]
                label.font.size = FONT_SIZES["body"]
                label.font.color.rgb = get_color("ink_800")
                text = para.add_run(phase["teacher_actions"])
                text.font.name = FONTS["body"]
                text.font.size = FONT_SIZES["body"]
                text.font.color.rgb = get_color("ink_700")
            if phase.get("student_actions"):
                para = doc.add_paragraph()
                label = para.add_run("Student Actions: ")
                label.bold = True
                label.font.name = FONTS["body"]
                label.font.size = FONT_SIZES["body"]
                label.font.color.rgb = get_color("ink_800")
                text = para.add_run(phase["student_actions"])
                text.font.name = FONTS["body"]
                text.font.size = FONT_SIZES["body"]
                text.font.color.rgb = get_color("ink_700")
            if phase.get("key_points"):
                _add_styled_paragraph(doc, "Key Points:", bold=True, color_key="ink_800")
                _add_bullet_list(doc, phase["key_points"])
            if phase.get("differentiation_notes"):
                para = doc.add_paragraph()
                label = para.add_run("Differentiation: ")
                label.bold = True
                label.font.name = FONTS["body"]
                label.font.size = FONT_SIZES["body"]
                label.font.color.rgb = get_color("gold_600")
                text = para.add_run(phase["differentiation_notes"])
                text.font.name = FONTS["body"]
                text.font.size = FONT_SIZES["body"]
                text.font.color.rgb = get_color("ink_600")
                text.italic = True
            doc.add_paragraph()

    # Exit Assessment
    exit_assess = session.get("exit_assessment", {})
    if exit_assess:
        _add_section_header(doc, "Exit Assessment", navy_accent)
        if exit_assess.get("type"):
            _add_key_value(doc, "Type", exit_assess["type"])
        if exit_assess.get("description"):
            _add_styled_paragraph(doc, exit_assess["description"])
        doc.add_paragraph()

    # Differentiation Overview (detailed) with color-coded levels
    if diff:
        _add_section_header(doc, "Differentiation Details", navy_accent)
        for level_key, level_data in diff.items():
            if isinstance(level_data, dict):
                level_name = LEVEL_NAMES.get(level_key, level_key.replace("_", " ").title())
                level_color = get_level_accent(level_key)

                # Level name with color
                para = doc.add_paragraph()
                level_run = para.add_run(f"■ {level_name}")
                level_run.bold = True
                level_run.font.name = FONTS["body"]
                level_run.font.size = FONT_SIZES["heading2"]
                level_run.font.color.rgb = level_color

                if level_data.get("focus"):
                    para = doc.add_paragraph()
                    label = para.add_run("Focus: ")
                    label.bold = True
                    label.font.name = FONTS["body"]
                    label.font.size = FONT_SIZES["body"]
                    label.font.color.rgb = get_color("ink_800")
                    text = para.add_run(level_data["focus"])
                    text.font.name = FONTS["body"]
                    text.font.size = FONT_SIZES["body"]
                    text.font.color.rgb = get_color("ink_700")
                if level_data.get("key_scaffolds"):
                    _add_styled_paragraph(doc, "Scaffolds:", bold=True, color_key="ink_800")
                    _add_bullet_list(doc, level_data["key_scaffolds"])
                if level_data.get("monitor_for"):
                    para = doc.add_paragraph()
                    label = para.add_run("Monitor for: ")
                    label.bold = True
                    label.font.name = FONTS["body"]
                    label.font.size = FONT_SIZES["body"]
                    label.font.color.rgb = get_color("ink_800")
                    text = para.add_run(level_data["monitor_for"])
                    text.font.name = FONTS["body"]
                    text.font.size = FONT_SIZES["body"]
                    text.font.color.rgb = get_color("ink_600")
                    text.italic = True
                doc.add_paragraph()

    # EL Support Summary
    el_support = teacher_guide.get("el_support_summary", {})
    if el_support:
        _add_section_header(doc, "English Learner Support", get_color("emerging"))
        for el_level, support_data in el_support.items():
            if isinstance(support_data, dict):
                para = doc.add_paragraph()
                level_run = para.add_run(f"■ {el_level.title()}")
                level_run.bold = True
                level_run.font.name = FONTS["body"]
                level_run.font.size = FONT_SIZES["heading2"]
                level_run.font.color.rgb = get_color("emerging")
                if support_data.get("key_vocabulary_to_preteach"):
                    _add_styled_paragraph(doc, "Pre-teach Vocabulary:", bold=True, color_key="ink_800")
                    _add_bullet_list(doc, support_data["key_vocabulary_to_preteach"])
                if support_data.get("visual_supports_needed"):
                    _add_styled_paragraph(doc, "Visual Supports:", bold=True, color_key="ink_800")
                    _add_bullet_list(doc, support_data["visual_supports_needed"])
                if support_data.get("partner_recommendations"):
                    para = doc.add_paragraph()
                    label = para.add_run("Partner Recommendations: ")
                    label.bold = True
                    label.font.name = FONTS["body"]
                    label.font.size = FONT_SIZES["body"]
                    label.font.color.rgb = get_color("ink_800")
                    text = para.add_run(support_data["partner_recommendations"])
                    text.font.name = FONTS["body"]
                    text.font.size = FONT_SIZES["body"]
                    text.font.color.rgb = get_color("ink_700")
        doc.add_paragraph()

    # Materials List
    materials = teacher_guide.get("materials_list", [])
    if materials:
        _add_section_header(doc, "Materials Needed", navy_accent)
        _add_bullet_list(doc, materials)
        doc.add_paragraph()

    # Common Misconceptions
    misconceptions = teacher_guide.get("common_misconceptions", [])
    if misconceptions:
        _add_section_header(doc, "Common Misconceptions", get_color("error"))
        for misc in misconceptions:
            if isinstance(misc, dict):
                para = doc.add_paragraph()
                label = para.add_run("Misconception: ")
                label.bold = True
                label.font.name = FONTS["body"]
                label.font.size = FONT_SIZES["body"]
                label.font.color.rgb = get_color("error")
                text = para.add_run(misc.get("misconception", ""))
                text.font.name = FONTS["body"]
                text.font.size = FONT_SIZES["body"]
                text.font.color.rgb = get_color("ink_700")

                para2 = doc.add_paragraph()
                label2 = para2.add_run("How to Address: ")
                label2.bold = True
                label2.font.name = FONTS["body"]
                label2.font.size = FONT_SIZES["body"]
                label2.font.color.rgb = get_color("success")
                text2 = para2.add_run(misc.get("how_to_address", ""))
                text2.font.name = FONTS["body"]
                text2.font.size = FONT_SIZES["body"]
                text2.font.color.rgb = get_color("ink_700")
            else:
                _add_bullet_list(doc, [str(misc)])
        doc.add_paragraph()

    # Discussion Prompts
    prompts = teacher_guide.get("discussion_prompts", [])
    if prompts:
        _add_section_header(doc, "Discussion Prompts")
        _add_numbered_list(doc, prompts)
        doc.add_paragraph()

    # Formative Assessment Ideas
    assessments = teacher_guide.get("formative_assessment_ideas", [])
    if assessments:
        _add_section_header(doc, "Formative Assessment Ideas")
        _add_bullet_list(doc, assessments)
        doc.add_paragraph()


def generate_student_material_section(
    doc: Document,
    level_key: str,
    level_data: dict,
    day_num: Optional[int] = None,
    lesson_title: str = ""
) -> None:
    """Generate a student material section for one level.

    Args:
        doc: Document to add to
        level_key: The level identifier (e.g., 'below_level', 'at_level')
        level_data: The student material data for this level
        day_num: Optional day number for multi-day lessons
        lesson_title: The lesson title for the header
    """
    level_name = LEVEL_NAMES.get(level_key, level_key.replace("_", " ").title())
    accent_color = get_level_accent(level_key)

    # Check for multi-day structure
    if "days" in level_data:
        # Multi-day format
        for day_data in level_data["days"]:
            day = day_data.get("day", 1)
            day_title = day_data.get("title", lesson_title or f"Day {day}")
            # Use student header instead of plain heading
            _add_student_header(doc, day_title, level_name, accent_color)
            _generate_student_day_content(doc, day_data, level_key, accent_color)
            doc.add_page_break()
    else:
        # Single-day format
        display_title = lesson_title or level_name
        if day_num:
            display_title = f"Day {day_num} - {display_title}"
        # Use student header with name/date fields
        _add_student_header(doc, display_title, level_name, accent_color)
        _generate_student_day_content(doc, level_data, level_key, accent_color)


def _generate_student_day_content(doc: Document, data: dict, level_key: str = None, accent_color: RGBColor = None) -> None:
    """Generate content for a single day of student materials.

    Args:
        doc: Document to add to
        data: The student material data
        level_key: The readiness level key (for color matching)
        accent_color: The accent color for this level
    """
    if accent_color is None:
        accent_color = get_level_accent(level_key) if level_key else get_color("navy_700")

    # Header / I CAN Statement
    header = data.get("header", {})
    if header:
        # Use goal box for I CAN statement
        i_can = header.get("i_can_statement") or header.get("student_objective")
        if i_can:
            _add_goal_box(doc, i_can, accent_color)

    # Vocabulary - render as styled table
    vocab = data.get("vocabulary", [])
    if vocab:
        _add_section_header(doc, "Vocabulary", accent_color)

        # Create vocabulary table
        table = doc.add_table(rows=1 + len(vocab), cols=2)
        table.style = 'Table Grid'
        table.autofit = False
        table.columns[0].width = Inches(2.0)
        table.columns[1].width = Inches(5.0)

        # Header row
        accent_hex = f"{accent_color.red:02x}{accent_color.green:02x}{accent_color.blue:02x}" if hasattr(accent_color, 'red') else COLORS["navy_700"]
        for i, header_text in enumerate(["Term", "Definition"]):
            cell = table.rows[0].cells[i]
            cell.text = header_text
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.name = FONTS["body"]
                run.font.size = FONT_SIZES["body"]
                run.font.color.rgb = get_color("white")
            set_cell_shading(cell, accent_hex)

        # Data rows
        for row_idx, word in enumerate(vocab):
            if isinstance(word, dict):
                term = word.get("term", "")
                definition = word.get("definition", "")
                example = word.get("example", "")
                full_def = f"{definition}\nExample: {example}" if example else definition
            else:
                term = str(word)
                full_def = ""

            term_cell = table.rows[row_idx + 1].cells[0]
            term_cell.text = term
            for run in term_cell.paragraphs[0].runs:
                run.bold = True
                run.font.name = FONTS["body"]
                run.font.size = FONT_SIZES["body"]
                run.font.color.rgb = get_color("ink_800")

            def_cell = table.rows[row_idx + 1].cells[1]
            def_cell.text = full_def
            for run in def_cell.paragraphs[0].runs:
                run.font.name = FONTS["body"]
                run.font.size = FONT_SIZES["body"]
                run.font.color.rgb = get_color("ink_700")

        doc.add_paragraph()

    # Worked Example
    worked = data.get("worked_example", {})
    if worked:
        _add_section_header(doc, "Worked Example", accent_color)
        if worked.get("problem"):
            para = doc.add_paragraph()
            problem_label = para.add_run("Problem: ")
            problem_label.bold = True
            problem_label.font.name = FONTS["body"]
            problem_label.font.size = FONT_SIZES["body"]
            problem_label.font.color.rgb = get_color("ink_800")
            problem_text = para.add_run(worked["problem"])
            problem_text.font.name = FONTS["body"]
            problem_text.font.size = FONT_SIZES["body"]
            problem_text.font.color.rgb = get_color("ink_700")

        steps = worked.get("steps", [])
        if steps:
            steps_para = doc.add_paragraph()
            steps_label = steps_para.add_run("Steps:")
            steps_label.bold = True
            steps_label.font.name = FONTS["body"]
            steps_label.font.size = FONT_SIZES["body"]

            for step in steps:
                if isinstance(step, dict):
                    step_num = step.get("step_number", "")
                    action = step.get("action", "")
                    result = step.get("result", "")
                    step_para = doc.add_paragraph(f"Step {step_num}: {action}", style='List Number')
                    for run in step_para.runs:
                        run.font.name = FONTS["body"]
                        run.font.size = FONT_SIZES["body"]
                    if result:
                        result_para = doc.add_paragraph()
                        result_label = result_para.add_run("→ ")
                        result_label.font.color.rgb = accent_color
                        result_text = result_para.add_run(result)
                        result_text.font.name = FONTS["body"]
                        result_text.font.size = FONT_SIZES["body"]
                        result_text.font.color.rgb = get_color("ink_600")
                        result_text.italic = True
                else:
                    doc.add_paragraph(str(step), style='List Number')

        # Handle solution_summary (used by at_level instead of detailed steps)
        if worked.get("solution_summary"):
            para = doc.add_paragraph()
            sol_label = para.add_run("Solution: ")
            sol_label.bold = True
            sol_label.font.name = FONTS["body"]
            sol_label.font.color.rgb = get_color("ink_800")
            sol_text = para.add_run(worked["solution_summary"])
            sol_text.font.name = FONTS["body"]
            sol_text.font.color.rgb = get_color("ink_700")

        # Handle final solution if present
        if worked.get("solution"):
            para = doc.add_paragraph()
            ans_label = para.add_run("Answer: ")
            ans_label.bold = True
            ans_label.font.name = FONTS["body"]
            ans_label.font.color.rgb = accent_color
            ans_text = para.add_run(worked["solution"])
            ans_text.font.name = FONTS["body"]
            ans_text.font.color.rgb = get_color("ink_800")
            ans_text.bold = True

        doc.add_paragraph()

    # Guided Practice
    guided = data.get("guided_practice", [])
    if guided:
        _add_section_header(doc, "Guided Practice", accent_color)
        for i, item in enumerate(guided, 1):
            if isinstance(item, dict):
                problem = item.get("problem", item.get("prompt", ""))
                prob_para = doc.add_paragraph()
                num_run = prob_para.add_run(f"{i}. ")
                num_run.bold = True
                num_run.font.name = FONTS["body"]
                num_run.font.size = FONT_SIZES["body"]
                num_run.font.color.rgb = accent_color
                prob_run = prob_para.add_run(problem)
                prob_run.font.name = FONTS["body"]
                prob_run.font.size = FONT_SIZES["body"]
                prob_run.font.color.rgb = get_color("ink_700")

                # Support both 'scaffold' (below level) and 'hint' (approaching level)
                hint = item.get("scaffold") or item.get("hint")
                if hint:
                    hint_para = doc.add_paragraph()
                    hint_label = hint_para.add_run("Hint: ")
                    hint_label.italic = True
                    hint_label.font.name = FONTS["body"]
                    hint_label.font.size = FONT_SIZES["hint"]
                    hint_label.font.color.rgb = get_color("ink_500")
                    hint_text = hint_para.add_run(hint)
                    hint_text.italic = True
                    hint_text.font.name = FONTS["body"]
                    hint_text.font.size = FONT_SIZES["hint"]
                    hint_text.font.color.rgb = get_color("ink_600")

                # Add workspace box only if explicitly requested
                if item.get("workspace"):
                    _add_workspace_box(doc, num_lines=3, accent_color=accent_color)
            else:
                prob_para = doc.add_paragraph()
                num_run = prob_para.add_run(f"{i}. ")
                num_run.bold = True
                num_run.font.color.rgb = accent_color
                prob_para.add_run(str(item))
        doc.add_paragraph()

    # Independent Practice
    independent = data.get("independent_practice", [])
    if independent:
        _add_section_header(doc, "Independent Practice", accent_color)
        for i, item in enumerate(independent, 1):
            if isinstance(item, dict):
                problem = item.get("problem", item.get("prompt", ""))
                prob_para = doc.add_paragraph()
                num_run = prob_para.add_run(f"{i}. ")
                num_run.bold = True
                num_run.font.name = FONTS["body"]
                num_run.font.size = FONT_SIZES["body"]
                num_run.font.color.rgb = accent_color
                prob_run = prob_para.add_run(problem)
                prob_run.font.name = FONTS["body"]
                prob_run.font.size = FONT_SIZES["body"]
                prob_run.font.color.rgb = get_color("ink_700")

                # Add workspace box only if explicitly requested
                if item.get("workspace"):
                    _add_workspace_box(doc, num_lines=3, accent_color=accent_color)
            else:
                prob_para = doc.add_paragraph()
                num_run = prob_para.add_run(f"{i}. ")
                num_run.bold = True
                num_run.font.color.rgb = accent_color
                prob_para.add_run(str(item))
        doc.add_paragraph()

    # Practice Problems (alternative structure)
    practice = data.get("practice_problems", [])
    if practice:
        _add_section_header(doc, "Practice Problems", accent_color)
        for i, item in enumerate(practice, 1):
            if isinstance(item, dict):
                problem = item.get("problem", "")
                prob_para = doc.add_paragraph()
                num_run = prob_para.add_run(f"{i}. ")
                num_run.bold = True
                num_run.font.name = FONTS["body"]
                num_run.font.size = FONT_SIZES["body"]
                num_run.font.color.rgb = accent_color
                prob_run = prob_para.add_run(problem)
                prob_run.font.name = FONTS["body"]
                prob_run.font.size = FONT_SIZES["body"]
                prob_run.font.color.rgb = get_color("ink_700")

                # Add workspace box only if explicitly requested
                if item.get("workspace"):
                    _add_workspace_box(doc, num_lines=3, accent_color=accent_color)
            else:
                prob_para = doc.add_paragraph()
                num_run = prob_para.add_run(f"{i}. ")
                num_run.bold = True
                num_run.font.color.rgb = accent_color
                prob_para.add_run(str(item))
        doc.add_paragraph()

    # Graphic Organizer - render as actual table
    organizer = data.get("graphic_organizer", {})
    if organizer:
        _add_section_header(doc, "Graphic Organizer", accent_color)
        _render_graphic_organizer(doc, organizer, accent_color)

    # Sentence Frames
    frames = data.get("sentence_frames", [])
    if frames:
        _add_section_header(doc, "Sentence Frames", accent_color)
        # Render as styled info box
        for frame in frames:
            _add_info_box(doc, "Frame", str(frame), accent_color)

    # Word Bank
    word_bank = data.get("word_bank", [])
    if word_bank:
        _add_section_header(doc, "Word Bank", accent_color)
        # Create a styled word bank box
        table = doc.add_table(rows=1, cols=1)
        table.autofit = False
        table.columns[0].width = Inches(7.0)
        cell = table.rows[0].cells[0]

        para = cell.paragraphs[0]
        words_text = "  •  ".join(word_bank)
        run = para.add_run(words_text)
        run.font.name = FONTS["body"]
        run.font.size = FONT_SIZES["body"]
        run.font.color.rgb = get_color("ink_700")
        run.bold = True
        para.paragraph_format.space_before = Pt(8)
        para.paragraph_format.space_after = Pt(8)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        set_cell_shading(cell, COLORS["ink_50"])
        accent_hex = f"{accent_color.red:02x}{accent_color.green:02x}{accent_color.blue:02x}" if hasattr(accent_color, 'red') else COLORS["ink_400"]
        set_cell_border(cell, "top", accent_hex, width=8, style="single")
        set_cell_border(cell, "bottom", accent_hex, width=8, style="single")
        set_cell_border(cell, "left", COLORS["ink_50"], width=0, style="nil")
        set_cell_border(cell, "right", COLORS["ink_50"], width=0, style="nil")

        doc.add_paragraph()

    # Application Problem (at_level)
    app_problem = data.get("application_problem", {})
    if app_problem:
        _add_section_header(doc, "Application Problem", accent_color)
        if app_problem.get("context"):
            context_para = doc.add_paragraph()
            context_run = context_para.add_run(app_problem["context"])
            context_run.font.name = FONTS["body"]
            context_run.font.size = FONT_SIZES["body"]
            context_run.font.color.rgb = get_color("ink_700")

        if app_problem.get("question"):
            para = doc.add_paragraph()
            q_label = para.add_run("Question: ")
            q_label.bold = True
            q_label.font.name = FONTS["body"]
            q_label.font.color.rgb = get_color("ink_800")
            q_text = para.add_run(app_problem["question"])
            q_text.font.name = FONTS["body"]
            q_text.font.color.rgb = get_color("ink_700")

        # Add workspace box instead of "[Work Space]"
        _add_workspace_box(doc, num_lines=4, accent_color=accent_color)

    # Extension Challenge
    extension = data.get("extension_challenge", {})
    if extension:
        _add_section_header(doc, "Extension Challenge", accent_color)
        if extension.get("title"):
            para = doc.add_paragraph()
            title_run = para.add_run(extension["title"])
            title_run.bold = True
            title_run.font.name = FONTS["body"]
            title_run.font.size = FONT_SIZES["body"]
            title_run.font.color.rgb = accent_color

        if extension.get("description"):
            desc_para = doc.add_paragraph()
            desc_run = desc_para.add_run(extension["description"])
            desc_run.font.name = FONTS["body"]
            desc_run.font.size = FONT_SIZES["body"]
            desc_run.font.color.rgb = get_color("ink_700")

        if extension.get("guiding_questions"):
            q_para = doc.add_paragraph()
            q_label = q_para.add_run("Guiding Questions:")
            q_label.bold = True
            q_label.font.name = FONTS["body"]
            q_label.font.size = FONT_SIZES["body"]
            _add_bullet_list(doc, extension["guiding_questions"], accent_color)

        _add_workspace_box(doc, num_lines=4, accent_color=accent_color)

    # Reflection
    reflection = data.get("reflection", {})
    if reflection:
        _add_section_header(doc, "Reflection", accent_color)
        if reflection.get("prompt"):
            prompt_para = doc.add_paragraph()
            prompt_run = prompt_para.add_run(reflection["prompt"])
            prompt_run.font.name = FONTS["body"]
            prompt_run.font.size = FONT_SIZES["body"]
            prompt_run.font.color.rgb = get_color("ink_700")
            prompt_run.italic = True

        # Add workspace for reflection response
        _add_workspace_box(doc, num_lines=5, accent_color=accent_color)


def generate_combined_document(curriculum: dict, include_udl: bool = False) -> Document:
    """Generate a single combined DOCX with all curriculum content.

    Structure:
    1. Teacher Guide (all days if multi-day)
    2. Student Materials - Below Level (all days)
    3. Student Materials - Approaching Level (all days)
    4. Student Materials - At Level (all days)
    5. Student Materials - Above Level (all days)
    """
    doc = Document()

    # Add page numbers to footer
    _add_page_numbers(doc)

    teacher_guide = curriculum.get("teacher_guide", {})
    student_materials = curriculum.get("student_materials", {})

    # Check if multi-day format (teacher guide has "days" array)
    is_multi_day = "days" in teacher_guide

    # ===== TEACHER GUIDE SECTION =====
    _add_styled_heading(doc, "Teacher Guide", level=1)

    if is_multi_day:
        # Multi-day teacher guide
        meta = teacher_guide.get("metadata", {})
        unit_overview = teacher_guide.get("unit_overview", {})

        # Unit title and overview
        doc.add_paragraph(meta.get("title", "Lesson Plan"))
        if unit_overview.get("learning_arc"):
            _add_key_value(doc, "Learning Arc", unit_overview["learning_arc"])
        if unit_overview.get("essential_questions"):
            doc.add_paragraph("Essential Questions:")
            _add_bullet_list(doc, unit_overview["essential_questions"])
        doc.add_paragraph()

        # Each day - merge top-level metadata with day-specific data
        top_meta = teacher_guide.get("metadata", {})
        for day_data in teacher_guide.get("days", []):
            day_num = day_data.get("day", 1)
            # Create merged data: day-specific metadata takes precedence, but inherit from top-level
            day_meta = day_data.get("metadata", {})
            merged_meta = {
                "title": day_data.get("title", day_meta.get("title", f"Day {day_num}")),
                "grade": day_meta.get("grade", top_meta.get("grade")),
                "subject": day_meta.get("subject", top_meta.get("subject")),
                "topic": day_meta.get("topic", top_meta.get("topic")),
                "duration_minutes": day_meta.get("duration_minutes", top_meta.get("duration_minutes_per_day", top_meta.get("duration_minutes"))),
                "standards_addressed": day_meta.get("standards_addressed", top_meta.get("standards_addressed", [])),
                "pedagogical_approach": day_meta.get("pedagogical_approach", top_meta.get("pedagogical_approach")),
            }
            # Create merged day data with proper metadata
            merged_day = {**day_data, "metadata": merged_meta}
            generate_teacher_guide_section(doc, merged_day, day_num=day_num)
            doc.add_page_break()

        # Shared sections (differentiation, EL support, etc.)
        diff = teacher_guide.get("differentiation_overview", {})
        if diff:
            _add_section_header(doc, "Differentiation Overview (All Days)")
            for level_key, level_data in diff.items():
                if isinstance(level_data, dict):
                    level_name = LEVEL_NAMES.get(level_key, level_key.replace("_", " ").title())
                    para = doc.add_paragraph()
                    para.add_run(level_name).bold = True
                    if level_data.get("focus"):
                        doc.add_paragraph(f"Focus: {level_data['focus']}")
                    if level_data.get("key_scaffolds"):
                        _add_bullet_list(doc, level_data["key_scaffolds"])
            doc.add_page_break()
    else:
        # Single-day teacher guide
        generate_teacher_guide_section(doc, teacher_guide)
        doc.add_page_break()

    # ===== STUDENT MATERIALS SECTIONS =====
    _add_styled_heading(doc, "Student Materials", level=1)

    # Add intro paragraph with styling
    intro_para = doc.add_paragraph()
    intro_run = intro_para.add_run("The following pages contain differentiated student handouts for all readiness levels.")
    intro_run.font.name = FONTS["body"]
    intro_run.font.size = FONT_SIZES["body"]
    intro_run.font.color.rgb = get_color("ink_600")
    intro_run.italic = True
    doc.add_page_break()

    # Get lesson title for student headers
    meta = teacher_guide.get("metadata", {})
    lesson_title = meta.get("title", "")

    # Generate each level
    levels_order = ["below_level", "approaching_level", "at_level", "above_level"]
    for level_key in levels_order:
        level_data = student_materials.get(level_key, {})
        if level_data:
            generate_student_material_section(doc, level_key, level_data, lesson_title=lesson_title)
            doc.add_page_break()

    return doc


def save_combined_document(curriculum: dict, output_path: str, include_udl: bool = False) -> str:
    """Generate and save a combined DOCX document.

    Args:
        curriculum: The curriculum dictionary from LLM
        output_path: Directory to save the file
        include_udl: Whether to include UDL documentation (future)

    Returns:
        Path to the saved file
    """
    doc = generate_combined_document(curriculum, include_udl)

    # Generate filename
    teacher_guide = curriculum.get("teacher_guide", {})
    meta = teacher_guide.get("metadata", {})

    title = meta.get("title", "Lesson")
    # Clean title for filename
    clean_title = "".join(c if c.isalnum() or c in " -_" else "" for c in title)
    clean_title = clean_title.replace(" ", "_")[:50]

    filename = f"{clean_title}_lesson_plan.docx"
    filepath = Path(output_path) / filename

    doc.save(str(filepath))
    return filename
